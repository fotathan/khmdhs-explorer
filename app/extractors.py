"""
extractors.py — file classification, zip unpacking, table extraction.

Each extracted table is a plain dict:
{
    "id": str,            # stable id used by the UI / exporter
    "source": str,        # original (display) file path, e.g. "ΠΑΡΑΡΤΗΜΑΤΑ.zip/Τεύχος.docx"
    "locator": str,       # where inside the file: "Sheet 'Φύλλο1'", "Page 3, table 2", ...
    "rows": list[list[str]],
    "n_rows": int,
    "n_cols": int,
}

Files that cannot be processed are reported as "skipped" with a reason,
so scanned PDFs / images are *flagged*, never silently dropped.
"""

from __future__ import annotations

import io
import os
import re
import uuid
import zipfile
from dataclasses import dataclass, field

import pdfplumber
from docx import Document as DocxDocument
from openpyxl import load_workbook

try:
    import xlrd  # legacy .xls
    HAS_XLRD = True
except ImportError:
    HAS_XLRD = False

MAX_ZIP_DEPTH = 3
MIN_TABLE_ROWS = 2
MIN_TABLE_COLS = 2

TABLE_FORMATS = {".xlsx", ".xlsm", ".xls", ".docx", ".pdf", ".csv"}
IMAGE_FORMATS = {".png", ".jpg", ".jpeg", ".tif", ".tiff", ".bmp", ".gif", ".webp"}


@dataclass
class FileReport:
    """Per-file outcome shown in the results screen."""
    source: str
    status: str            # "ok" | "no_tables" | "scanned" | "image" | "unsupported" | "error"
    detail: str = ""
    tables: list = field(default_factory=list)
    entry_id: str = ""      # FileEntry this report came from (enables per-file OCR)
    n_pages: int = 0        # page count, where known (shown on the OCR button)


@dataclass
class FileEntry:
    """A single concrete file discovered in an upload (possibly inside zips)."""
    id: str
    source: str            # display path, e.g. "ΠΑΡΑΡΤΗΜΑΤΑ.zip → Τεύχος.docx"
    name: str               # basename
    ext: str
    size: int
    data: bytes

    @property
    def kind(self) -> str:
        """Coarse classification used by the file-selection screen."""
        if self.ext in TABLE_FORMATS:
            return "tables"
        if self.ext in IMAGE_FORMATS:
            return "image"
        if self.ext in (".doc", ".rar"):
            return "unsupported"
        return "other"

    @property
    def size_human(self) -> str:
        n = float(self.size)
        for unit in ("B", "KB", "MB", "GB"):
            if n < 1024 or unit == "GB":
                return f"{n:.0f} {unit}" if unit == "B" else f"{n:.1f} {unit}"
            n /= 1024
        return f"{self.size} B"


# ---------------------------------------------------------------- helpers

def _ext(name: str) -> str:
    return os.path.splitext(name)[1].lower()


def _clean_cell(value) -> str:
    if value is None:
        return ""
    s = str(value)
    return re.sub(r"\s+", " ", s).strip()


def _trim_table(rows: list[list[str]]) -> list[list[str]]:
    """Drop fully-empty leading/trailing rows and columns."""
    rows = [r for r in rows if any(c for c in r)]
    if not rows:
        return []
    width = max(len(r) for r in rows)
    rows = [r + [""] * (width - len(r)) for r in rows]
    # trim empty columns from both ends
    cols_with_data = [i for i in range(width) if any(r[i] for r in rows)]
    if not cols_with_data:
        return []
    lo, hi = cols_with_data[0], cols_with_data[-1] + 1
    return [r[lo:hi] for r in rows]


def _make_table(source: str, locator: str, rows: list[list[str]]) -> dict | None:
    rows = _trim_table(rows)
    if len(rows) < MIN_TABLE_ROWS:
        return None
    if max(len(r) for r in rows) < MIN_TABLE_COLS:
        return None
    return {
        "id": uuid.uuid4().hex[:12],
        "source": source,
        "locator": locator,
        "rows": rows,
        "n_rows": len(rows),
        "n_cols": len(rows[0]),
    }


def _fix_zip_name(info: zipfile.ZipInfo) -> str:
    """
    Greek zips made on Windows often store filenames in CP737/CP1253 without
    the UTF-8 flag; the zipfile module then mis-decodes them as CP437.
    Re-decode in that case so 'Διακήρυξη.pdf' doesn't come out as mojibake.
    """
    if info.flag_bits & 0x800:          # UTF-8 flag set, name is fine
        return info.filename
    raw = info.filename.encode("cp437", errors="replace")
    for enc in ("utf-8", "cp737", "cp1253"):
        try:
            return raw.decode(enc)
        except UnicodeDecodeError:
            continue
    return info.filename


# ---------------------------------------------------------------- per-format extractors

def _extract_xlsx(data: bytes, source: str) -> FileReport:
    tables = []
    wb = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    try:
        for ws in wb.worksheets:
            rows = [[_clean_cell(c) for c in row] for row in ws.iter_rows(values_only=True)]
            t = _make_table(source, f"Sheet “{ws.title}”", rows)
            if t:
                tables.append(t)
    finally:
        wb.close()
    if not tables:
        return FileReport(source, "no_tables", "No usable data found in any sheet.")
    return FileReport(source, "ok", tables=tables)


def _extract_xls(data: bytes, source: str) -> FileReport:
    if not HAS_XLRD:
        return FileReport(source, "unsupported", "Legacy .xls support requires the 'xlrd' package.")
    tables = []
    book = xlrd.open_workbook(file_contents=data)
    for sheet in book.sheets():
        rows = [[_clean_cell(sheet.cell_value(r, c)) for c in range(sheet.ncols)]
                for r in range(sheet.nrows)]
        t = _make_table(source, f"Sheet “{sheet.name}”", rows)
        if t:
            tables.append(t)
    if not tables:
        return FileReport(source, "no_tables", "No usable data found in any sheet.")
    return FileReport(source, "ok", tables=tables)


def _extract_csv(data: bytes, source: str) -> FileReport:
    import csv
    text = None
    for enc in ("utf-8-sig", "utf-8", "cp1253", "iso-8859-7"):
        try:
            text = data.decode(enc)
            break
        except UnicodeDecodeError:
            continue
    if text is None:
        return FileReport(source, "error", "Could not decode CSV text.")
    try:
        dialect = csv.Sniffer().sniff(text[:4096], delimiters=",;\t|")
    except csv.Error:
        dialect = csv.excel
    rows = [[_clean_cell(c) for c in row] for row in csv.reader(io.StringIO(text), dialect)]
    t = _make_table(source, "CSV data", rows)
    if not t:
        return FileReport(source, "no_tables", "No usable data found.")
    return FileReport(source, "ok", tables=[t])


def _extract_docx(data: bytes, source: str) -> FileReport:
    doc = DocxDocument(io.BytesIO(data))
    tables = []
    for i, tbl in enumerate(doc.tables, start=1):
        rows = [[_clean_cell(cell.text) for cell in row.cells] for row in tbl.rows]
        t = _make_table(source, f"Table {i}", rows)
        if t:
            tables.append(t)
    if not tables:
        return FileReport(source, "no_tables", "Document contains no tables.")
    return FileReport(source, "ok", tables=tables)


def _norm_row(row: list[str]) -> tuple:
    return tuple(re.sub(r"\s+", "", c).casefold() for c in row)


def _column_edges(tbl) -> list[int]:
    """Left x-coordinates of a pdfplumber table's columns, rounded."""
    xs = sorted({round(cell[0]) for cell in tbl.cells if cell})
    return xs


def _edges_match(a: list[int], b: list[int], tol: int = 10) -> bool:
    return len(a) == len(b) and all(abs(x - y) <= tol for x, y in zip(a, b))


def _width(rows: list[list[str]]) -> int:
    return max((len(r) for r in rows), default=0)


def _continues(group: list[dict], nxt: dict) -> bool:
    """Does fragment `nxt` continue the table run in `group`?"""
    prev = group[-1]
    if nxt["page"] != prev["page"] + 1:
        return False
    if prev["idx"] != prev["count_on_page"] or nxt["idx"] != 1:
        return False  # must be last table on its page → first table on the next
    if not prev["rows"] or not nxt["rows"]:
        return False
    if _width(prev["rows"]) != _width(nxt["rows"]):
        return False
    header = group[0]["rows"][0]
    return (_edges_match(prev["edges"], nxt["edges"])
            or _norm_row(nxt["rows"][0]) == _norm_row(header))


def compress_pages(pages) -> str:
    """[1,2,3,7,9,10] -> '1–3, 7, 9–10' (for UI summaries and provenance)."""
    pages = sorted(set(pages))
    if not pages:
        return ""
    runs, start, prev = [], pages[0], pages[0]
    for p in pages[1:]:
        if p == prev + 1:
            prev = p
            continue
        runs.append((start, prev))
        start = prev = p
    runs.append((start, prev))
    return ", ".join(f"{a}" if a == b else f"{a}–{b}" for a, b in runs)


def _extract_pdf(data: bytes, source: str,
                 pages: set[int] | None = None) -> FileReport:
    fragments: list[dict] = []
    pages_total = 0
    pages_with_text = 0

    with pdfplumber.open(io.BytesIO(data)) as pdf:
        for page_no, page in enumerate(pdf.pages, start=1):
            if pages is not None and page_no not in pages:
                continue
            pages_total += 1
            text = (page.extract_text() or "").strip()
            if len(text) >= 20:
                pages_with_text += 1
            found = page.find_tables()
            for ti, tbl in enumerate(found, start=1):
                raw = tbl.extract()
                rows = [[_clean_cell(c) for c in row] for row in raw]
                fragments.append({
                    "page": page_no,
                    "idx": ti,
                    "count_on_page": len(found),
                    "rows": rows,
                    "edges": _column_edges(tbl),
                })

    # group consecutive-page fragments that look like one continuing table
    groups: list[list[dict]] = []
    for frag in fragments:
        if groups and _continues(groups[-1], frag):
            groups[-1].append(frag)
        else:
            groups.append([frag])

    tables: list[dict] = []
    for group in groups:
        if len(group) == 1:
            f = group[0]
            t = _make_table(source, f"Page {f['page']}, table {f['idx']}", f["rows"])
            if t:
                tables.append(t)
            continue

        # stitch: keep the first fragment whole, drop repeated headers afterwards
        header = _norm_row(group[0]["rows"][0])
        stitched_rows: list[list[str]] = list(group[0]["rows"])
        for f in group[1:]:
            rows = f["rows"]
            if rows and _norm_row(rows[0]) == header:
                rows = rows[1:]
            stitched_rows.extend(rows)

        p1, p2 = group[0]["page"], group[-1]["page"]
        stitched = _make_table(
            source, f"Pages {p1}–{p2} (stitched from {len(group)} pages)", stitched_rows
        )
        if stitched:
            stitched["role"] = "stitched"
            tables.append(stitched)
        for k, f in enumerate(group, start=1):
            t = _make_table(
                source, f"Page {f['page']}, table {f['idx']} (part {k}/{len(group)})",
                f["rows"],
            )
            if t:
                if stitched:
                    t["role"] = "fragment"
                    t["group"] = stitched["id"]
                tables.append(t)

    if tables:
        return FileReport(source, "ok", tables=tables)
    subset = f" (checked pages {compress_pages(pages)})" if pages else ""
    if pages_total and pages_with_text / pages_total < 0.2:
        return FileReport(
            source, "scanned",
            f"{pages_total - pages_with_text} of {pages_total} checked pages have no text layer — "
            f"this looks like a scanned document{subset}. Use the OCR button to extract its tables.",
            n_pages=pages_total,
        )
    return FileReport(source, "no_tables",
                      f"Text layer present, but no tables were detected{subset}.",
                      n_pages=pages_total)


# ---------------------------------------------------------------- entry points

# ---------------------------------------------------------------- plain-text extraction
#
# Separate from the table extractors above: this pulls the readable *text* out
# of a file (for a full-text field / search), not its tables. Used both by the
# KHMDHS ingester (auto fill on import) and the manual edit-page button.
#
# Each `_text_*` returns a string, or "" when the format carries no extractable
# text. `extract_text_from_entry` returns None when there is no text layer at
# all (e.g. a scanned PDF or a bare image) so callers can cleanly distinguish
# "empty document" from "needs OCR" and skip it in automatic mode.

# A PDF whose total extracted text is below this many non-space characters is
# treated as having no real text layer (scanned). Tunable; deliberately small.
MIN_PDF_TEXT_CHARS = 20


def _text_pdf(data: bytes, pages: set[int] | None = None) -> str:
    out: list[str] = []
    with pdfplumber.open(io.BytesIO(data)) as pdf:
        for i, page in enumerate(pdf.pages, start=1):
            if pages is not None and i not in pages:
                continue
            txt = page.extract_text() or ""
            if txt.strip():
                out.append(txt)
    return "\n\n".join(out).strip()


def _text_docx(data: bytes) -> str:
    doc = DocxDocument(io.BytesIO(data))
    parts: list[str] = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
    # include table cell text too — tender τεύχη put a lot of content in tables
    for tbl in doc.tables:
        for row in tbl.rows:
            cells = [_clean_cell(c.text) for c in row.cells]
            line = " | ".join(c for c in cells if c)
            if line:
                parts.append(line)
    return "\n".join(parts).strip()


def _text_xlsx(data: bytes) -> str:
    wb = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    parts: list[str] = []
    for ws in wb.worksheets:
        parts.append(f"[{ws.title}]")
        for row in ws.iter_rows(values_only=True):
            cells = [_clean_cell(c) for c in row]
            line = " | ".join(c for c in cells if c)
            if line:
                parts.append(line)
    return "\n".join(parts).strip()


def _text_xls(data: bytes) -> str:
    if xlrd is None:
        return ""
    book = xlrd.open_workbook(file_contents=data)
    parts: list[str] = []
    for sh in book.sheets():
        parts.append(f"[{sh.name}]")
        for r in range(sh.nrows):
            cells = [_clean_cell(sh.cell_value(r, c)) for c in range(sh.ncols)]
            line = " | ".join(c for c in cells if c)
            if line:
                parts.append(line)
    return "\n".join(parts).strip()


def _text_csv(data: bytes) -> str:
    for enc in ("utf-8-sig", "utf-8", "cp1253", "iso-8859-7", "latin-1"):
        try:
            return data.decode(enc).strip()
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace").strip()


def extract_text_from_entry(entry: FileEntry,
                            pages: set[int] | None = None) -> str | None:
    """Pull plain text out of one collected file.

    Returns the text, or None when there is no extractable text layer (scanned
    PDF, image, or unsupported type) — so automatic callers can skip it and
    leave OCR to a deliberate manual step. `pages` (PDF only) restricts which
    1-based pages are read.
    """
    ext = entry.ext
    try:
        if ext == ".pdf":
            txt = _text_pdf(entry.data, pages)
            return txt if len(txt.replace(" ", "")) >= MIN_PDF_TEXT_CHARS else None
        if ext == ".docx":
            return _text_docx(entry.data) or None
        if ext in (".xlsx", ".xlsm"):
            return _text_xlsx(entry.data) or None
        if ext == ".xls":
            return _text_xls(entry.data) or None
        if ext == ".csv":
            return _text_csv(entry.data) or None
        # images, .doc, .rar, anything else: no text layer here
        return None
    except Exception:  # noqa: BLE001 — never let one bad file break a batch
        return None


def extract_text_from_upload(filename: str, data: bytes) -> str | None:
    """Convenience for the ingester: take a raw downloaded attachment (which may
    itself be a zip of τεύχη), unpack it, extract text from every text-bearing
    file inside, and return the concatenation — or None if nothing yielded text.
    """
    entries, _errors = collect_files(filename, data)
    chunks: list[str] = []
    for e in entries:
        t = extract_text_from_entry(e)
        if t:
            # label each file so a multi-doc attachment stays legible
            chunks.append(f"=== {e.source} ===\n{t}")
    combined = "\n\n".join(chunks).strip()
    return combined or None


def collect_files(filename: str, data: bytes, depth: int = 0,
                  prefix: str = "") -> tuple[list[FileEntry], list[FileReport]]:
    """
    Phase 1: walk one upload (recursing into zips) and return every concrete
    file found, without parsing any of them yet. Bad archives come back as
    error FileReports so they still surface in the results screen.
    """
    source = f"{prefix}{filename}"
    ext = _ext(filename)

    if ext == ".zip":
        if depth >= MAX_ZIP_DEPTH:
            return [], [FileReport(source, "error", "Zip nesting too deep.")]
        entries: list[FileEntry] = []
        errors: list[FileReport] = []
        try:
            with zipfile.ZipFile(io.BytesIO(data)) as zf:
                for info in zf.infolist():
                    if info.is_dir():
                        continue
                    inner_name = _fix_zip_name(info)
                    base = os.path.basename(inner_name)
                    if base.startswith((".", "~$")) or "__MACOSX" in inner_name:
                        continue
                    sub_entries, sub_errors = collect_files(
                        base, zf.read(info), depth + 1, prefix=f"{source} → "
                    )
                    entries.extend(sub_entries)
                    errors.extend(sub_errors)
        except zipfile.BadZipFile:
            return [], [FileReport(source, "error", "Not a valid zip archive.")]
        if not entries and not errors:
            errors.append(FileReport(source, "no_tables", "Archive is empty."))
        return entries, errors

    return [FileEntry(
        id=uuid.uuid4().hex[:12],
        source=source,
        name=filename,
        ext=ext,
        size=len(data),
        data=data,
    )], []


def extract_entry(entry: FileEntry, pages: set[int] | None = None) -> FileReport:
    """Phase 2: parse one collected file and extract its tables.

    `pages` (PDF only): restrict parsing to these 1-based page numbers."""
    source, data, ext = entry.source, entry.data, entry.ext
    try:
        if ext in (".xlsx", ".xlsm"):
            return _extract_xlsx(data, source)
        if ext == ".xls":
            return _extract_xls(data, source)
        if ext == ".csv":
            return _extract_csv(data, source)
        if ext == ".docx":
            return _extract_docx(data, source)
        if ext == ".pdf":
            return _extract_pdf(data, source, pages)
        if ext in IMAGE_FORMATS:
            return FileReport(source, "image",
                              "Image file — use the OCR button to extract any tables it contains.",
                              n_pages=1)
        if ext == ".doc":
            return FileReport(source, "unsupported",
                              "Legacy .doc format — re-save as .docx, or open in Word/LibreOffice.")
        if ext == ".rar":
            return FileReport(source, "unsupported",
                              ".rar archives are not supported — extract manually and upload the contents.")
        return FileReport(source, "unsupported", f"Unsupported file type ({ext or 'no extension'}).")
    except Exception as exc:  # noqa: BLE001 — surface, don't crash the batch
        return FileReport(source, "error", f"Failed to process: {exc.__class__.__name__}: {exc}")
